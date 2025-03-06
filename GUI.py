import tkinter as tk
from tkinter import ttk, messagebox, StringVar
from tkinter import *
import mysql.connector
import mysql
from ttkthemes import ThemedTk
from PIL import Image, ImageTk
from tkinter import filedialog
from docx import Document
import warnings
warnings.warn("This feature is pending deprecation", PendingDeprecationWarning)


class PatientManagementApp:
    def __init__(self, root):
        self.root = root
        self.root.title("J F K Patient Management System")
        self.root.geometry("1800x1000")
        self.root.resizable(True, False)

        self.apply_styles()

        # Background Image
        bg_image = Image.open("background.jpg").resize((2500, 1800))
        self.bg_photo = ImageTk.PhotoImage(bg_image)
        bg_label = tk.Label(self.root, image=self.bg_photo)
        bg_label.place(x=0, y=0, relwidth=1, relheight=1)

        # Sidebar
        self.sidebar = tk.Frame(self.root, bg="#2d3e50")
        self.sidebar.place(x=0, y=0, width=250, height=700)

        sidebar_buttons = [
            ("Home", self.show_home),
            ("Register", self.patient_register),
            ("Patients Login", self.patient_login),
            ("Admin login", self.admin_dashboard),
            ("About Us", self.show_about),
            ("Exit", self.root.quit),
        ]

        for idx, (text, command) in enumerate(sidebar_buttons):
            ttk.Button(
                self.sidebar,
                text=text,
                command=command,
                style="Sidebar.TButton",
            ).pack(pady=5, fill="x")

        # Main Content Area
        self.content_frame = tk.Frame(self.root, bg="#f0f4fa")
        self.content_frame.place(x=250, y=0, width=1010, height=700)

        # Load Home Page
        self.show_home()

    def apply_styles(self):
        style = ttk.Style()
        style.theme_use("clam")

        # Sidebar Button Style
        style.configure(
            "Sidebar.TButton",
            font=("Helvetica", 12, "bold"),
            foreground="#ffffff",
            background="#2d3e50",
            padding=10,
        )
        style.map(
            "Sidebar.TButton",
            background=[("active", "#3c78b4")],
            foreground=[("active", "#ffffff")],
        )

        # Main Button Style
        style.configure(
            "Main.TButton",
            font=("Helvetica", 12, "bold"),
            foreground="#ffffff",
            background="#5a9bd5",
            padding=10,
        )
        style.map(
            "Main.TButton",
            background=[("active", "#3c78b4")],
            foreground=[("active", "#ffffff")],
        )

    def show_home(self):
        for widget in self.content_frame.winfo_children():
            widget.destroy()

        tk.Label(
            self.content_frame,
            text="Welcome to the J.F.K Patient Management System!",
            font=("Helvetica", 18, "bold"),
            bg="#f0f4fa",
        ).pack(pady=20)

        tk.Label(
            self.content_frame,
            text="Here your records are manage in a very secure way\n Patients can view appointments,medical records, and more.",
            font=("Helvetica", 14),
            bg="#f0f4fa",
        ).pack(pady=10)

        ttk.Button(
            self.content_frame,
            text="Register Today",
            command=self.patient_register,
            style="Main.TButton",
        ).pack(pady=20)

    def add_navigation_buttons(self, back_command, show_logout=False):
        """Adds Back and optionally Logout buttons to the content frame."""
        button_frame = tk.Frame(self.content_frame, bg="#f0f4fa")
        button_frame.pack(pady=20, side=tk.BOTTOM, fill=tk.X)

        # Back Button
        ttk.Button(
            button_frame,
            text="Back",
            command=back_command,
            style="Main.TButton",
        ).pack(side=tk.LEFT, padx=10)

    def patient_register(self):
        for widget in self.content_frame.winfo_children():
            widget.destroy()

        tk.Label(
            self.content_frame,
            text="Patient Registration",
            font=("Helvetica", 16, "bold"),
            bg="#f0f4fa",
        ).pack(pady=10)

        # Patient Registration Form
        form_frame = tk.Frame(self.content_frame, bg="#f0f4fa")
        form_frame.pack(pady=20)

        fields = [
            ("First Name:", "first_name"),
            ("Last Name:", "last_name"),
            ("Date of Birth (YYYY-MM-DD):", "dob"),
            ("Email:", "email"),
            ("Password:", "password"),
            ("Contact:", "contact"),
            ("Address:", "address"),
        ]

        entries = {}
        for i, (label_text, field_key) in enumerate(fields):
            tk.Label(form_frame, text=label_text, bg="#f0f4fa").grid(row=i, column=0, pady=5, sticky="w")
            entry = ttk.Entry(form_frame, width=30, show="*" if field_key == "password" else None)
            entry.grid(row=i, column=1, pady=5)
            entries[field_key] = entry

        def submit_registration():
            data = {key: entry.get() for key, entry in entries.items()}
            if all(data.values()):
                try:
                    self.save_to_database(data)
                    messagebox.showinfo("Registration Success", "Patient Registered Successfully!")
                    self.show_home()
                except Exception as e:
                    messagebox.showerror("Database Error", str(e))
            else:
                messagebox.showerror("Error", "All fields are required!")

        ttk.Button(form_frame, text="Submit", command=submit_registration, style="Main.TButton").grid(row=len(fields),
                                                                                                      column=0,
                                                                                                      columnspan=2,
                                                                                                      pady=20)
        self.add_navigation_buttons(back_command=self.show_home, show_logout=True)


    def save_to_database(self, data):
        """Save patient data to the database."""
        try:
            conn = mysql.connector.connect(
                host="sql5.freesqldatabase.com",
                user="sql5765928",
                password="dhu72jpWAU",
                database="sql5765928"
            )
            cursor = conn.cursor()
            query = """
            INSERT INTO patients (first_name, last_name, dob, email, password, contact, address)
            VALUES (%s, %s, %s, %s, %s, %s, %s)
            """
            cursor.execute(query, (
            data["first_name"], data["last_name"], data["dob"], data["email"], data["password"], data["contact"],
            data["address"]))
            conn.commit()
            cursor.close()
            conn.close()
        except mysql.connector.Error as err:
            raise Exception(f"Error: {err}")

    def patient_login(self):
        # Clear the content frame
        for widget in self.content_frame.winfo_children():
            widget.destroy()

        # Add the "Patient Login" title
        tk.Label(
            self.content_frame,
            text="Patient Login",
            font=("Helvetica", 16, "bold"),
            bg="#f0f4fa",
        ).pack(pady=10)

        # Create a form frame
        form_frame = tk.Frame(self.content_frame, bg="#f0f4fa")
        form_frame.pack(pady=20)

        # Email field
        tk.Label(form_frame, text="Email:", bg="#f0f4fa").grid(row=0, column=0, pady=5, sticky="w")
        email_entry = ttk.Entry(form_frame, width=30)
        email_entry.grid(row=0, column=1, pady=5)

        # Password field
        tk.Label(form_frame, text="Password:", bg="#f0f4fa").grid(row=1, column=0, pady=5, sticky="w")
        password_entry = ttk.Entry(form_frame, show="*", width=30)
        password_entry.grid(row=1, column=1, pady=5)

        # Submit button logic
        def submit_login():
            email = email_entry.get().strip()
            password = password_entry.get().strip()

            # Input validation: Check if fields are empty
            if not email or not password:
                messagebox.showerror("Validation Error", "All fields are required!")
                return

            try:
                # Connect to the external MySQL database
                conn = mysql.connector.connect(
                    host="sql5.freesqldatabase.com",
                    user="sql5765928",
                    password="dhu72jpWAU",
                    database="sql5765928"     # Database name
                )
                cursor = conn.cursor(dictionary=True)

                # Query to check if the patient exists
                query = "SELECT * FROM patients WHERE email = %s AND password = %s"
                cursor.execute(query, (email, password))
                result = cursor.fetchone()

                # Close the database connection
                cursor.close()
                conn.close()

                if result:
                    # If login is successful, show the patient dashboard
                    self.show_patient_dashboard(result)
                else:
                    # If login fails, show an error message
                    messagebox.showerror("Error", "Invalid email or password!")

            except mysql.connector.Error as err:
                # Handle database connection or query errors
                messagebox.showerror("Database Error", f"Error: {err}")

        # Add the login button
        ttk.Button(form_frame, text="Login", command=submit_login, style="Main.TButton").grid(
            row=2, column=0, columnspan=2, pady=20
        )

        # Add navigation buttons
        self.add_navigation_buttons(back_command=self.show_home, show_logout=True)

    def admin_dashboard(self):
        """Admin login window."""
        # Clear content_frame
        for widget in self.content_frame.winfo_children():
            widget.destroy()

        # Title Label
        tk.Label(self.content_frame, text="Admin Login", font=("Helvetica", 24, "bold"), bg="#f0f4fa").pack(pady=50)

        # Username Entry
        self.username_label = tk.Label(self.content_frame, text="Username", font=("Helvetica", 16), bg="#f0f4fa")
        self.username_label.pack(pady=5)
        self.username_entry = ttk.Entry(self.content_frame, font=("Helvetica", 16), width=30)
        self.username_entry.pack(pady=10)

        # Password Entry
        self.password_label = tk.Label(self.content_frame, text="Password", font=("Helvetica", 16), bg="#f0f4fa")
        self.password_label.pack(pady=5)
        self.password_entry = ttk.Entry(self.content_frame, font=("Helvetica", 16), width=30, show="*")
        self.password_entry.pack(pady=10)

        # Login Button
        ttk.Button(self.content_frame, text="Login", command=self.verify_admin_login, style="TButton").pack(pady=20)

        # Back Button
        ttk.Button(self.content_frame, text="Back", command=self.show_home, style="TButton").pack(pady=10)

    def verify_admin_login(self):
        """Verify admin login credentials."""
        username = self.username_entry.get()
        password = self.password_entry.get()

        if not username or not password:
            messagebox.showerror("Error", "Please fill out both fields.")
            return

        try:
            with mysql.connector.connect(
                    host="sql5.freesqldatabase.com",
                    user="sql5765928",
                    password="dhu72jpWAU",
                    database="sql5765928"
            ) as conn:
                cursor = conn.cursor(dictionary=True)
                query = "SELECT * FROM admin WHERE username = %s AND password = %s"
                cursor.execute(query, (username, password))
                admin_data = cursor.fetchone()

            if admin_data:
                messagebox.showinfo("Success", "Login Successful!")
                self.show_admin_dashboard(admin_data)
            else:
                messagebox.showerror("Error", "Invalid username or password.")

        except mysql.connector.Error as err:
            messagebox.showerror("Database Error", f"Error: {err}")

    def show_admin_dashboard(self, admin_data=None):
        """Display the admin dashboard after successful login."""
        if admin_data:
            self.admin_data = admin_data  # Store admin_data for later use

        # Clear any existing widgets in the content frame
        for widget in self.content_frame.winfo_children():
            widget.destroy()

        # Admin Welcome Message
        tk.Label(
            self.content_frame,
            text=f"Welcome, {self.admin_data['username']}!",
            font=("Helvetica", 24, "bold"),
            bg="#f0f4fa"
        ).pack(pady=50)

        # Dashboard Options
        options_frame = tk.Frame(self.content_frame, bg="#f0f4fa")
        options_frame.pack(pady=20, fill=tk.X)

        ttk.Button(options_frame, text="View Patients", command=self.view_patients, style="Main.TButton").pack(pady=10)
        ttk.Button(options_frame, text="Manage Patients", command=self.manage_patients, style="Main.TButton").pack(
            pady=10)
        ttk.Button(options_frame, text="Manage Appointments", command=self.manage_appointments,
                   style="Main.TButton").pack(pady=10)
        ttk.Button(options_frame, text="Manage Medical Records", command=self.manage_medical_records,
                   style="Main.TButton").pack(pady=10)

        # Add a button to generate the report
        generate_report_button = tk.Button(self.content_frame, text="Generate Report", command=self.generate_report,
                                           font=("Arial", 12))
        generate_report_button.pack(pady=20)

        # Logout button
        ttk.Button(options_frame, text="Logout", command=self.show_home, style="Main.TButton").pack(pady=10)

    def generate_report(self):
        """Function to generate a report containing data from specific tables in the database."""
        try:
            # Connect to the database
            with mysql.connector.connect(
                    host="sql5.freesqldatabase.com",
                    user="sql5765928",
                    password="dhu72jpWAU",
                    database="sql5765928"
            ) as conn:
                cursor = conn.cursor(dictionary=True)

                # Create a new Word document
                doc = Document()
                doc.add_heading('Database Report', 0)

                # List of specific tables to include in the report
                tables = ['admin', 'patients', 'medicalrecords', 'appointments']

                # Loop through the tables and add their data to the report
                for table_name in tables:
                    doc.add_heading(f'Table: {table_name.capitalize()}', level=1)

                    # Get data from the table
                    cursor.execute(f"SELECT * FROM {table_name}")
                    rows = cursor.fetchall()

                    # Add the table data to the document if data exists
                    if rows:
                        # Add a table to the document: First row for headers
                        columns = [desc[0] for desc in cursor.description]
                        table = doc.add_table(rows=1, cols=len(columns))

                        # Add column headers to the first row
                        hdr_cells = table.rows[0].cells
                        for i, column in enumerate(columns):
                            hdr_cells[i].text = column

                        # Add data rows
                        for row in rows:
                            row_cells = table.add_row().cells
                            for i, column in enumerate(columns):
                                row_cells[i].text = str(row[column])

                    else:
                        doc.add_paragraph("No data available.")

                # Ask the admin for a location to save the report
                file_path = filedialog.asksaveasfilename(defaultextension=".docx",
                                                         filetypes=[("Word Documents", "*.docx")])

                if file_path:
                    # Save the document at the chosen location
                    doc.save(file_path)
                    messagebox.showinfo("Success", f"Report saved successfully at {file_path}.")
                else:
                    messagebox.showwarning("No File Selected",
                                           "No file location was selected. The report was not saved.")
        except mysql.connector.Error as err:
            messagebox.showerror("Database Error", f"Error: {err}")
    def view_patients(self):
        """View all patients with an enhanced GUI and scrollbars."""
        # Clear previous content
        for widget in self.content_frame.winfo_children():
            widget.destroy()

        try:
            # Connect to the database
            with mysql.connector.connect(
                    host="sql5.freesqldatabase.com",
                    user="sql5765928",
                    password="dhu72jpWAU",
                    database="sql5765928"
            ) as conn:
                cursor = conn.cursor(dictionary=True)
                query = "SELECT * FROM patients"
                cursor.execute(query)
                patients = cursor.fetchall()

            # Add a title label
            title = tk.Label(self.content_frame, text="Patient Records", font=("Arial", 16, "bold"), fg="blue")
            title.pack(pady=10)

            # Create a style for the Treeview with smaller font
            style = ttk.Style()
            style.configure("Treeview", font=("Arial", 10), rowheight=20)
            style.configure("Treeview.Heading", font=("Arial", 12, "bold"))

            # Create the Treeview with multiple columns
            columns = ("Patient ID", "Name", "Gender", "DOB", "Phone", "Email", "Address", "Created At")
            tree = ttk.Treeview(self.content_frame, columns=columns, show="headings")

            # Define column headings
            for col in columns:
                tree.heading(col, text=col)
                tree.column(col, anchor="center", width=150)

            # Add data to the Treeview
            for patient in patients:
                tree.insert("", "end", values=(
                    patient["patient_id"],
                    f"{patient['first_name']} {patient['last_name']}",
                    patient["gender"],
                    patient["dob"],
                    patient["contact"],
                    patient["email"],
                    patient["address"],
                    patient["created_at"]
                ))

            # Add vertical and horizontal scrollbars
            v_scrollbar = ttk.Scrollbar(self.content_frame, orient="vertical", command=tree.yview)
            h_scrollbar = ttk.Scrollbar(self.content_frame, orient="horizontal", command=tree.xview)
            tree.configure(yscrollcommand=v_scrollbar.set, xscrollcommand=h_scrollbar.set)

            # Pack the scrollbars
            v_scrollbar.pack(side="right", fill="y")
            h_scrollbar.pack(side="bottom", fill="x")

            # Pack the Treeview
            tree.pack(padx=10, pady=10, expand=True, fill="both")

        except mysql.connector.Error as err:
            messagebox.showerror("Database Error", f"Error: {err}")

        # Add navigation buttons
        self.add_navigation_buttons(back_command=lambda: self.show_admin_dashboard())

    def manage_patients(self):
        """Manage patient information (Add, Update, Delete)."""
        for widget in self.content_frame.winfo_children():
            widget.destroy()

        manage_frame = tk.Frame(self.content_frame)
        manage_frame.pack(pady=20)
        ttk.Label=Label(manage_frame,text="Choose your choice of operation when it comes to managing patient:",font=("Arial",18))
        ttk.Label.pack()


        ttk.Button(manage_frame, text="Add Patient", command=self.add_patient, style="Main.TButton").pack(pady=10)
        ttk.Button(manage_frame, text="Update Patient", command=self.update_patient, style="Main.TButton").pack(pady=10)
        ttk.Button(manage_frame, text="Delete Patient", command=self.delete_patient, style="Main.TButton").pack(pady=10)
        self.add_navigation_buttons(back_command=lambda: self.show_admin_dashboard())

    def add_patient(self):
        """Function to add a new patient in the same window with a modern and clean design."""
        # Clear previous content in the content_frame
        for widget in self.content_frame.winfo_children():
            widget.destroy()

        # Define labels and input fields for patient details
        fields = [
            "First Name", "Last Name", "Gender", "Date of Birth (YYYY-MM-DD)",
            "Contact", "Email", "Password", "Address"
        ]
        input_vars = {field: tk.StringVar() for field in fields}

        # Create a form layout in the content_frame with a modern design
        # Header Label
        header_label = tk.Label(
            self.content_frame,
            text="Add New Patient",
            font=("Arial", 20, "bold"),
            fg="#2C3E50",  # Dark blue color
            bg="#F0F4FA"  # Light background
        )
        header_label.grid(row=0, column=0, columnspan=2, pady=20, sticky="ew")

        # Create a frame for the form fields
        form_frame = tk.Frame(self.content_frame, bg="#F0F4FA", padx=20, pady=20)
        form_frame.grid(row=1, column=0, columnspan=2, sticky="nsew")

        # Add form fields with labels and entry widgets
        for idx, field in enumerate(fields):
            # Label
            label = tk.Label(
                form_frame,
                text=field,
                font=("Arial", 12),
                bg="#F0F4FA",
                fg="#34495E",  # Dark gray color
                anchor="w"
            )
            label.grid(row=idx, column=0, padx=10, pady=5, sticky="w")

            # Entry Field
            entry = tk.Entry(
                form_frame,
                textvariable=input_vars[field],
                font=("Arial", 12),
                bg="white",
                fg="#2C3E50",  # Dark blue color
                relief="flat",
                bd=2
            )
            entry.grid(row=idx, column=1, padx=10, pady=5, sticky="ew")

        # Function to save the patient to the database
        def save_patient():
            # Collect data from input fields
            patient_data = {field: input_vars[field].get() for field in fields}
            patient_data['created_at'] = patient_data['updated_at'] = 'CURRENT_TIMESTAMP'

            # Validate required fields
            missing_fields = [field for field, value in patient_data.items() if not value]
            if missing_fields:
                messagebox.showwarning(
                    "Validation Error",
                    f"Please fill out the following required fields: {', '.join(missing_fields)}"
                )
                return

            try:
                # Insert data into the database
                with mysql.connector.connect(
                        host="sql5.freesqldatabase.com",
                        user="sql5765928",
                        password="dhu72jpWAU",
                        database="sql5765928"
                ) as conn:
                    cursor = conn.cursor()
                    query = """
                        INSERT INTO patients (first_name, last_name, gender, dob, contact, email, password, address, created_at, updated_at)
                        VALUES (%s, %s, %s, %s, %s, %s, %s, %s, CURRENT_TIMESTAMP, CURRENT_TIMESTAMP)
                    """
                    cursor.execute(query, (
                        patient_data["First Name"], patient_data["Last Name"], patient_data["Gender"],
                        patient_data["Date of Birth (YYYY-MM-DD)"], patient_data["Contact"],
                        patient_data["Email"], patient_data["Password"], patient_data["Address"]
                    ))
                    conn.commit()

                # Show success message and clear the form
                messagebox.showinfo("Success", "Patient added successfully!")
                for var in input_vars.values():
                    var.set("")

            except mysql.connector.Error as err:
                messagebox.showerror("Database Error", f"Error: {err}")

        # Create a frame for buttons
        button_frame = tk.Frame(self.content_frame, bg="#F0F4FA", padx=20, pady=20)
        button_frame.grid(row=2, column=0, columnspan=2, sticky="ew")

        # Save Button
        save_button = tk.Button(
            button_frame,
            text="Save",
            command=save_patient,
            font=("Arial", 12, "bold"),
            bg="#27AE60",  # Green color
            fg="white",
            relief="flat",
            padx=20,
            pady=10
        )
        save_button.pack(side="left", padx=10, pady=10)

        # Back Button
        back_button = tk.Button(
            button_frame,
            text="Back",
            command=self.manage_patients,
            font=("Arial", 12, "bold"),
            bg="#E74C3C",  # Red color
            fg="white",
            relief="flat",
            padx=20,
            pady=10
        )
        back_button.pack(side="right", padx=10, pady=10)

        # Configure grid weights for responsive layout
        self.content_frame.grid_columnconfigure(0, weight=1)
        self.content_frame.grid_columnconfigure(1, weight=1)
        self.content_frame.grid_rowconfigure(1, weight=1)

    def update_patient(self):
        """Function to update patient information in the same window with a modern and clean design."""
        # Clear previous content in the content_frame
        for widget in self.content_frame.winfo_children():
            widget.destroy()

        # Header with a modern design
        header_frame = tk.Frame(self.content_frame, bg="#4F6D7A", padx=20, pady=20)
        header_frame.pack(fill="x", pady=10)

        Label(header_frame, text="Update Patient Information", font=("Arial", 18, "bold"), bg="#4F6D7A",
              fg="white").pack()

        # Define labels and input fields for patient details
        fields = [
            "Patient ID (Required)", "First Name", "Last Name", "Gender", "Date of Birth (YYYY-MM-DD)",
            "Contact", "Email", "Password", "Address"
        ]
        input_vars = {field: tk.StringVar() for field in fields}

        # Create a frame for the form
        form_frame = tk.Frame(self.content_frame, bg="#f0f4fa", padx=20, pady=20)
        form_frame.pack(fill="both", expand=True)

        # Add labels and entry fields with modern styling
        for idx, field in enumerate(fields):
            label = tk.Label(form_frame, text=field, font=("Arial", 12), bg="#f0f4fa", fg="#34495e")
            label.grid(row=idx, column=0, padx=10, pady=5, sticky="w")

            entry = tk.Entry(form_frame, textvariable=input_vars[field], font=("Arial", 12), width=30, bd=2,
                             relief="groove")
            entry.grid(row=idx, column=1, padx=10, pady=5)

        # Function to update the patient record
        def save_updates():
            # Collect data from input fields
            patient_data = {field: input_vars[field].get().strip() for field in fields}

            # Ensure Patient ID is provided
            if not patient_data["Patient ID (Required)"]:
                messagebox.showwarning("Validation Error", "Patient ID is required to update the record!")
                return

            # Check if at least one field (other than Patient ID) is provided for update
            updatable_fields = [
                "First Name", "Last Name", "Gender", "Date of Birth (YYYY-MM-DD)", "Contact", "Email", "Password",
                "Address"
            ]
            if not any(patient_data[field] for field in updatable_fields):
                messagebox.showwarning("Validation Error", "At least one field must be filled to update the record!")
                return

            try:
                # Establish database connection
                with mysql.connector.connect(
                        host="sql5.freesqldatabase.com",
                        user="sql5765928",
                        password="dhu72jpWAU",
                        database="sql5765928"
                ) as conn:
                    cursor = conn.cursor()

                    # Dynamically create query for provided fields
                    updates = []
                    values = []
                    for field in updatable_fields:
                        if patient_data[field]:
                            column_name = field.lower().replace(" ", "_")
                            updates.append(f"{column_name} = %s")
                            values.append(patient_data[field])

                    # Include Patient ID for the WHERE clause
                    values.append(patient_data["Patient ID (Required)"])

                    # Build the query
                    update_query = f"""
                        UPDATE patients 
                        SET {', '.join(updates)}, updated_at = CURRENT_TIMESTAMP
                        WHERE patient_id = %s
                    """
                    cursor.execute(update_query, values)
                    conn.commit()

                    if cursor.rowcount > 0:
                        messagebox.showinfo("Success", "Patient information updated successfully!")
                    else:
                        messagebox.showwarning("Update Failed", "No patient found with the provided ID.")

            except mysql.connector.Error as err:
                messagebox.showerror("Database Error", f"Error: {err}")

        # Button frame for Save and Back buttons
        button_frame = tk.Frame(self.content_frame, bg="#f0f4fa", padx=20, pady=20)
        button_frame.pack(fill="x", pady=10)

        # Save button with modern styling
        save_button = tk.Button(button_frame, text="Save", command=save_updates, font=("Arial", 12, "bold"),
                                bg="#4CAF50", fg="white", padx=20, pady=10, relief="flat")
        save_button.pack(side="left", padx=10)

        # Back button with modern styling
        back_button = tk.Button(button_frame, text="Back", command=self.manage_patients, font=("Arial", 12, "bold"),
                                bg="#FF6347", fg="white", padx=20, pady=10, relief="flat")
        back_button.pack(side="right", padx=10)

        # Add hover effects for buttons
        def on_enter(event):
            event.widget.config(bg="#45a049" if event.widget == save_button else "#FF4500")

        def on_leave(event):
            event.widget.config(bg="#4CAF50" if event.widget == save_button else "#FF6347")

        save_button.bind("<Enter>", on_enter)
        save_button.bind("<Leave>", on_leave)
        back_button.bind("<Enter>", on_enter)
        back_button.bind("<Leave>", on_leave)

    def delete_patient(self):
        """Function to view detailed patient information and delete a selected record."""
        # Clear previous content in the content_frame
        for widget in self.content_frame.winfo_children():
            widget.destroy()

        # Instructions for the admin
        Label(self.content_frame, text="Delete Patient Record", font=("Arial", 18, "bold"), fg="red").grid(
            row=0, column=0, columnspan=3, pady=10)

        # Frame for the Treeview and Scrollbars
        table_frame = ttk.Frame(self.content_frame)
        table_frame.grid(row=1, column=0, columnspan=3, padx=10, pady=10, sticky="nsew")

        # Define a Treeview to display patient records with a further reduced height
        tree = ttk.Treeview(
            table_frame,
            columns=("Patient ID", "First Name", "Last Name", "Gender", "DOB", "Contact", "Email", "Address"),
            show="headings",
            height=5  # Further reduced height for a smaller table
        )
        tree.heading("Patient ID", text="Patient ID")
        tree.heading("First Name", text="First Name")
        tree.heading("Last Name", text="Last Name")
        tree.heading("Gender", text="Gender")
        tree.heading("DOB", text="Date of Birth")
        tree.heading("Contact", text="Contact")
        tree.heading("Email", text="Email")
        tree.heading("Address", text="Address")

        # Adjust column widths further to make the table more compact
        tree.column("Patient ID", width=60, anchor="center", stretch=True)  # Reduced width
        tree.column("First Name", width=100, stretch=True)  # Reduced width
        tree.column("Last Name", width=100, stretch=True)  # Reduced width
        tree.column("Gender", width=50, anchor="center", stretch=True)  # Reduced width
        tree.column("DOB", width=80, anchor="center", stretch=True)  # Reduced width
        tree.column("Contact", width=80, anchor="center", stretch=True)  # Reduced width
        tree.column("Email", width=130, stretch=True)  # Reduced width
        tree.column("Address", width=150, stretch=True)  # Reduced width

        # Add Scrollbars
        vertical_scrollbar = ttk.Scrollbar(table_frame, orient="vertical", command=tree.yview)
        horizontal_scrollbar = ttk.Scrollbar(table_frame, orient="horizontal", command=tree.xview)
        tree.configure(yscrollcommand=vertical_scrollbar.set, xscrollcommand=horizontal_scrollbar.set)

        # Pack scrollbars and treeview
        vertical_scrollbar.pack(side="right", fill="y")
        horizontal_scrollbar.pack(side="bottom", fill="x")
        tree.pack(expand=True, fill="both")

        # Function to load patient records into the table
        def load_patients():
            try:
                with mysql.connector.connect(
                        host="sql5.freesqldatabase.com",
                        user="sql5765928",
                        password="dhu72jpWAU",
                        database="sql5765928"
                ) as conn:
                    cursor = conn.cursor(dictionary=True)
                    query = """
                        SELECT patient_id, first_name, last_name, gender, dob, contact, email, address
                        FROM patients
                    """
                    cursor.execute(query)
                    patients = cursor.fetchall()

                # Check if records exist
                if not patients:
                    messagebox.showinfo("No Records", "No patient records found.")
                    return

                # Insert patient records into the treeview
                for patient in patients:
                    tree.insert("", "end", values=(
                        patient["patient_id"],
                        patient["first_name"],
                        patient["last_name"],
                        patient["gender"],
                        patient["dob"],
                        patient["contact"],
                        patient["email"],
                        patient["address"]
                    ))
            except mysql.connector.Error as err:
                messagebox.showerror("Database Error", f"Error: {err}")

        # Function to delete the selected patient
        def confirm_delete():
            selected_item = tree.selection()
            if not selected_item:
                messagebox.showwarning("Selection Error", "Please select a patient to delete.")
                return

            # Get the Patient ID of the selected record
            patient_id = tree.item(selected_item[0], "values")[0]

            # Confirm deletion
            confirm = messagebox.askyesno("Confirm Deletion",
                                          f"Are you sure you want to delete Patient ID {patient_id}?")
            if confirm:
                try:
                    with mysql.connector.connect(
                            host="sql5.freesqldatabase.com",
                            user="sql5765928",
                            password="dhu72jpWAU",
                            database="sql5765928"
                    ) as conn:
                        cursor = conn.cursor()
                        query = "DELETE FROM patients WHERE patient_id = %s"
                        cursor.execute(query, (patient_id,))
                        conn.commit()

                    if cursor.rowcount > 0:
                        messagebox.showinfo("Success", f"Patient ID {patient_id} has been deleted.")
                        tree.delete(selected_item[0])
                    else:
                        messagebox.showwarning("Deletion Failed", f"No patient found with ID {patient_id}.")
                except mysql.connector.Error as err:
                    messagebox.showerror("Database Error", f"Error: {err}")

        # Add Delete and Cancel buttons
        Button(self.content_frame, text="Delete", command=confirm_delete, font=("Arial", 12), bg="red",
               fg="white").grid(
            row=2, column=0, padx=10, pady=20, sticky="e")


        # Load patient records into the table
        load_patients()
        nav_frame = tk.Frame(self.content_frame)
        nav_frame.grid(row=2, column=1, columnspan=3, pady=20)

        # Add navigation buttons
        ttk.Button(nav_frame, text="Back", command=self.manage_patients).pack(pady=10)

    def manage_appointments(self):
        """Manage appointments (Create, Update, Delete)."""
        for widget in self.content_frame.winfo_children():
            widget.destroy()

        manage_frame = tk.Frame(self.content_frame)
        manage_frame.pack(pady=20)
        ttk.Label = Label(manage_frame, text="Choose your choice of operation when it comes to managing Appointments:",
                          font=("Arial", 18))
        ttk.Label.pack()

        ttk.Button(manage_frame, text="Create Appointment", command=self.create_appointment, style="Main.TButton").pack(
            pady=10)
        ttk.Button(manage_frame, text="Update Appointment", command=self.update_appointment,
                   style="Main.TButton").pack(pady=10)
        ttk.Button(manage_frame, text="Delete Appointment", command=self.delete_appointment,
                   style="Main.TButton").pack(pady=10)
        self.add_navigation_buttons(back_command=lambda: self.show_admin_dashboard())


    def create_appointment(self):
        """Function to create an appointment."""
        # Clear previous content in the content_frame
        for widget in self.content_frame.winfo_children():
            widget.destroy()

        # Instructions for the admin
        Label(self.content_frame, text="Create Appointment", font=("Arial", 18, "bold"), fg="blue").grid(
            row=0, column=0, columnspan=3, pady=10)

        # Frame for form fields
        form_frame = ttk.Frame(self.content_frame)
        form_frame.grid(row=1, column=0, columnspan=3, padx=10, pady=10)

        # Patient ID Input
        Label(form_frame, text="Patient ID:", font=("Arial", 12)).grid(row=0, column=0, padx=5, pady=5, sticky="e")
        patient_id_entry = Entry(form_frame, font=("Arial", 12))
        patient_id_entry.grid(row=0, column=1, padx=5, pady=5)

        # Doctor ID Input
        Label(form_frame, text="Doctor ID:", font=("Arial", 12)).grid(row=1, column=0, padx=5, pady=5, sticky="e")
        doctor_id_entry = Entry(form_frame, font=("Arial", 12))
        doctor_id_entry.grid(row=1, column=1, padx=5, pady=5)

        # Appointment Date Input (Ensure proper date format)
        Label(form_frame, text="Appointment Date (YYYY-MM-DD):", font=("Arial", 12)).grid(row=2, column=0, padx=5,
                                                                                          pady=5, sticky="e")
        appointment_date_entry = Entry(form_frame, font=("Arial", 12))
        appointment_date_entry.grid(row=2, column=1, padx=5, pady=5)

        # Status Input (e.g., Pending, Completed)
        Label(form_frame, text="Status:", font=("Arial", 12)).grid(row=3, column=0, padx=5, pady=5, sticky="e")
        status_entry = Entry(form_frame, font=("Arial", 12))
        status_entry.grid(row=3, column=1, padx=5, pady=5)

        # Purpose of Appointment Input
        Label(form_frame, text="Purpose:", font=("Arial", 12)).grid(row=4, column=0, padx=5, pady=5, sticky="e")
        purpose_entry = Entry(form_frame, font=("Arial", 12))
        purpose_entry.grid(row=4, column=1, padx=5, pady=5)

        # Function to submit the appointment creation
        def submit_appointment():
            # Get the values from the input fields
            patient_id = patient_id_entry.get()
            doctor_id = doctor_id_entry.get()
            appointment_date = appointment_date_entry.get()
            status = status_entry.get()
            purpose = purpose_entry.get()

            # Validate the inputs (e.g., check if all fields are filled)
            if not patient_id or not doctor_id or not appointment_date or not status or not purpose:
                messagebox.showwarning("Input Error", "Please fill in all fields.")
                return

            # You can add more validation here for specific fields, like date format

            try:
                # Connect to the database and insert the appointment
                with mysql.connector.connect(
                        host="sql5.freesqldatabase.com",
                        user="sql5765928",
                        password="dhu72jpWAU",
                        database="sql5765928"
                ) as conn:
                    cursor = conn.cursor()
                    query = """
                        INSERT INTO appointments (patient_id, doctor_id, appointment_date, status, purpose)
                        VALUES (%s, %s, %s, %s, %s)
                    """
                    cursor.execute(query, (patient_id, doctor_id, appointment_date, status, purpose))
                    conn.commit()

                # Show success message
                messagebox.showinfo("Success", "Appointment created successfully!")
            except mysql.connector.Error as err:
                messagebox.showerror("Database Error", f"Error: {err}")

        # Add Submit and Cancel buttons
        Button(self.content_frame, text="Submit", command=submit_appointment, font=("Arial", 12), bg="green",
               fg="white").grid(
            row=5, column=0, padx=10, pady=20, sticky="e")

        nav_frame = tk.Frame(self.content_frame)
        nav_frame.grid(row=5, column=1, columnspan=3, pady=20)

        # Add navigation buttons
        ttk.Button(nav_frame, text="Back", command=self.manage_appointments).pack(pady=10)

    def update_appointment(self):
        """Function to update an existing appointment."""
        # Clear previous content in the content_frame
        for widget in self.content_frame.winfo_children():
            widget.destroy()

        # Instructions for the admin
        Label(self.content_frame, text="Update Appointment", font=("Arial", 18, "bold"), fg="blue").grid(
            row=0, column=0, columnspan=3, pady=10)

        # Frame for the Treeview to display all appointments
        tree_frame = ttk.Frame(self.content_frame)
        tree_frame.grid(row=1, column=0, columnspan=3, padx=10, pady=10, sticky="nsew")

        # Define a Treeview to display all appointments
        tree = ttk.Treeview(
            tree_frame,
            columns=("Appointment ID", "Patient ID", "Doctor ID", "Appointment Date", "Status", "Purpose"),
            show="headings",
            height=8
        )
        tree.heading("Appointment ID", text="Appointment ID")
        tree.heading("Patient ID", text="Patient ID")
        tree.heading("Doctor ID", text="Doctor ID")
        tree.heading("Appointment Date", text="Appointment Date")
        tree.heading("Status", text="Status")
        tree.heading("Purpose", text="Purpose")

        tree.column("Appointment ID", width=100, anchor="center", stretch=True)
        tree.column("Patient ID", width=100, anchor="center", stretch=True)
        tree.column("Doctor ID", width=100, anchor="center", stretch=True)
        tree.column("Appointment Date", width=150, anchor="center", stretch=True)
        tree.column("Status", width=100, anchor="center", stretch=True)
        tree.column("Purpose", width=200, stretch=True)

        tree.pack(expand=True, fill="both")

        # Add Scrollbars
        vertical_scrollbar = ttk.Scrollbar(tree_frame, orient="vertical", command=tree.yview)
        horizontal_scrollbar = ttk.Scrollbar(tree_frame, orient="horizontal", command=tree.xview)
        tree.configure(yscrollcommand=vertical_scrollbar.set, xscrollcommand=horizontal_scrollbar.set)

        vertical_scrollbar.pack(side="right", fill="y")
        horizontal_scrollbar.pack(side="bottom", fill="x")

        # Function to load all appointments into the Treeview
        def load_appointments():
            try:
                # Connect to the database to fetch all appointments
                with mysql.connector.connect(
                        host="sql5.freesqldatabase.com",
                        user="sql5765928",
                        password="dhu72jpWAU",
                        database="sql5765928"
                ) as conn:
                    cursor = conn.cursor(dictionary=True)
                    query = "SELECT * FROM appointments"
                    cursor.execute(query)
                    appointments = cursor.fetchall()

                # Insert all appointments into the treeview
                for appointment in appointments:
                    tree.insert("", "end", values=(
                        appointment["appointment_id"],
                        appointment["patient_id"],
                        appointment["doctor_id"],
                        appointment["appointment_date"],
                        appointment["status"],
                        appointment["purpose"]
                    ))

            except mysql.connector.Error as err:
                messagebox.showerror("Database Error", f"Error: {err}")

        load_appointments()  # Load appointments when the screen is shown

        # Frame for form fields (for updating a selected appointment)
        form_frame = ttk.Frame(self.content_frame)
        form_frame.grid(row=2, column=0, columnspan=3, padx=10, pady=10)

        # Appointment ID Input (to be populated when a row is selected from the Treeview)
        Label(form_frame, text="Appointment ID:", font=("Arial", 12)).grid(row=0, column=0, padx=5, pady=5, sticky="e")
        appointment_id_entry = Entry(form_frame, font=("Arial", 12))
        appointment_id_entry.grid(row=0, column=1, padx=5, pady=5)

        # Patient ID Input
        Label(form_frame, text="Patient ID:", font=("Arial", 12)).grid(row=1, column=0, padx=5, pady=5, sticky="e")
        patient_id_entry = Entry(form_frame, font=("Arial", 12))
        patient_id_entry.grid(row=1, column=1, padx=5, pady=5)

        # Doctor ID Input
        Label(form_frame, text="Doctor ID:", font=("Arial", 12)).grid(row=2, column=0, padx=5, pady=5, sticky="e")
        doctor_id_entry = Entry(form_frame, font=("Arial", 12))
        doctor_id_entry.grid(row=2, column=1, padx=5, pady=5)

        # Appointment Date Input
        Label(form_frame, text="Appointment Date (YYYY-MM-DD):", font=("Arial", 12)).grid(row=3, column=0, padx=5,
                                                                                          pady=5, sticky="e")
        appointment_date_entry = Entry(form_frame, font=("Arial", 12))
        appointment_date_entry.grid(row=3, column=1, padx=5, pady=5)

        # Status Input
        Label(form_frame, text="Status:", font=("Arial", 12)).grid(row=4, column=0, padx=5, pady=5, sticky="e")
        status_entry = Entry(form_frame, font=("Arial", 12))
        status_entry.grid(row=4, column=1, padx=5, pady=5)

        # Purpose Input
        Label(form_frame, text="Purpose:", font=("Arial", 12)).grid(row=5, column=0, padx=5, pady=5, sticky="e")
        purpose_entry = Entry(form_frame, font=("Arial", 12))
        purpose_entry.grid(row=5, column=1, padx=5, pady=5)

        # Function to fetch selected appointment data and populate the form fields
        def fetch_selected_appointment():
            selected_item = tree.selection()
            if not selected_item:
                messagebox.showwarning("Selection Error", "Please select an appointment to update.")
                return

            appointment_id = tree.item(selected_item[0], "values")[0]

            try:
                # Fetch the selected appointment's details from the database
                with mysql.connector.connect(
                        host="sql5.freesqldatabase.com",
                        user="sql5765928",
                        password="dhu72jpWAU",
                        database="sql5765928"
                ) as conn:
                    cursor = conn.cursor(dictionary=True)
                    query = "SELECT * FROM appointments WHERE appointment_id = %s"
                    cursor.execute(query, (appointment_id,))
                    appointment = cursor.fetchone()

                if appointment:
                    # Populate the form fields with the selected appointment's data
                    appointment_id_entry.delete(0, tk.END)
                    appointment_id_entry.insert(0, appointment["appointment_id"])

                    patient_id_entry.delete(0, tk.END)
                    patient_id_entry.insert(0, appointment["patient_id"])

                    doctor_id_entry.delete(0, tk.END)
                    doctor_id_entry.insert(0, appointment["doctor_id"])

                    appointment_date_entry.delete(0, tk.END)
                    appointment_date_entry.insert(0, appointment["appointment_date"])

                    status_entry.delete(0, tk.END)
                    status_entry.insert(0, appointment["status"])

                    purpose_entry.delete(0, tk.END)
                    purpose_entry.insert(0, appointment["purpose"])

            except mysql.connector.Error as err:
                messagebox.showerror("Database Error", f"Error: {err}")

        # Button to fetch selected appointment data
        Button(self.content_frame, text="Fetch Appointment Data", command=fetch_selected_appointment,
               font=("Arial", 12), bg="blue", fg="white").grid(
            row=6, column=0, padx=10, pady=20, sticky="e")

        # Function to update the selected appointment
        def submit_update():
            appointment_id = appointment_id_entry.get()
            patient_id = patient_id_entry.get()
            doctor_id = doctor_id_entry.get()
            appointment_date = appointment_date_entry.get()
            status = status_entry.get()
            purpose = purpose_entry.get()

            # Validate the inputs
            if not appointment_id or not patient_id or not doctor_id or not appointment_date or not status or not purpose:
                messagebox.showwarning("Input Error", "Please fill in all fields.")
                return

            try:
                # Connect to the database and update the appointment
                with mysql.connector.connect(
                        host="sql5.freesqldatabase.com",
                        user="sql5765928",
                        password="dhu72jpWAU",
                        database="sql5765928"
                ) as conn:
                    cursor = conn.cursor()
                    query = """
                        UPDATE appointments 
                        SET patient_id = %s, doctor_id = %s, appointment_date = %s, status = %s, purpose = %s
                        WHERE appointment_id = %s
                    """
                    cursor.execute(query, (patient_id, doctor_id, appointment_date, status, purpose, appointment_id))
                    conn.commit()

                messagebox.showinfo("Success", "Appointment updated successfully!")
                # Refresh the appointments table
                tree.delete(*tree.get_children())  # Clear existing entries
                load_appointments()  # Reload appointments

            except mysql.connector.Error as err:
                messagebox.showerror("Database Error", f"Error: {err}")

        # Button to submit the updated data
        Button(self.content_frame, text="Update Appointment", command=submit_update, font=("Arial", 12), bg="green",
               fg="white").grid(
            row=6, column=1, columnspan=3, pady=10)
        nav_frame = tk.Frame(self.content_frame)
        nav_frame.grid(row=7, column=0, columnspan=3, pady=20)

        # Add navigation buttons
        ttk.Button(nav_frame, text="Back", command=self.manage_appointments).pack(pady=10)

    def delete_appointment(self):
        """Function to delete an appointment."""
        # Clear previous content in the content_frame
        for widget in self.content_frame.winfo_children():
            widget.destroy()

        # Instructions for the admin
        Label(self.content_frame, text="Delete Appointment", font=("Arial", 18, "bold"), fg="blue").grid(
            row=0, column=0, columnspan=3, pady=10)

        # Frame for the Treeview to display all appointments
        tree_frame = ttk.Frame(self.content_frame)
        tree_frame.grid(row=1, column=0, columnspan=3, padx=10, pady=10, sticky="nsew")

        # Define a Treeview to display all appointments
        tree = ttk.Treeview(
            tree_frame,
            columns=("Appointment ID", "Patient ID", "Doctor ID", "Appointment Date", "Status", "Purpose"),
            show="headings",
            height=8
        )
        tree.heading("Appointment ID", text="Appointment ID")
        tree.heading("Patient ID", text="Patient ID")
        tree.heading("Doctor ID", text="Doctor ID")
        tree.heading("Appointment Date", text="Appointment Date")
        tree.heading("Status", text="Status")
        tree.heading("Purpose", text="Purpose")

        tree.column("Appointment ID", width=100, anchor="center", stretch=True)
        tree.column("Patient ID", width=100, anchor="center", stretch=True)
        tree.column("Doctor ID", width=100, anchor="center", stretch=True)
        tree.column("Appointment Date", width=150, anchor="center", stretch=True)
        tree.column("Status", width=100, anchor="center", stretch=True)
        tree.column("Purpose", width=200, stretch=True)

        tree.pack(expand=True, fill="both")

        # Add Scrollbars
        vertical_scrollbar = ttk.Scrollbar(tree_frame, orient="vertical", command=tree.yview)
        horizontal_scrollbar = ttk.Scrollbar(tree_frame, orient="horizontal", command=tree.xview)
        tree.configure(yscrollcommand=vertical_scrollbar.set, xscrollcommand=horizontal_scrollbar.set)

        vertical_scrollbar.pack(side="right", fill="y")
        horizontal_scrollbar.pack(side="bottom", fill="x")

        # Function to load all appointments into the Treeview
        def load_appointments():
            try:
                # Connect to the database to fetch all appointments
                with mysql.connector.connect(
                        host="sql5.freesqldatabase.com",
                        user="sql5765928",
                        password="dhu72jpWAU",
                        database="sql5765928"
                ) as conn:
                    cursor = conn.cursor(dictionary=True)
                    query = "SELECT * FROM appointments"
                    cursor.execute(query)
                    appointments = cursor.fetchall()

                # Insert all appointments into the treeview
                for appointment in appointments:
                    tree.insert("", "end", values=(
                        appointment["appointment_id"],
                        appointment["patient_id"],
                        appointment["doctor_id"],
                        appointment["appointment_date"],
                        appointment["status"],
                        appointment["purpose"]
                    ))

            except mysql.connector.Error as err:
                messagebox.showerror("Database Error", f"Error: {err}")

        load_appointments()  # Load appointments when the screen is shown

        # Function to delete the selected appointment
        def delete_selected_appointment():
            selected_item = tree.selection()
            if not selected_item:
                messagebox.showwarning("Selection Error", "Please select an appointment to delete.")
                return

            appointment_id = tree.item(selected_item[0], "values")[0]

            # Ask for confirmation before deleting
            confirm = messagebox.askyesno("Confirm Deletion",
                                          f"Are you sure you want to delete appointment ID {appointment_id}?")
            if confirm:
                try:
                    # Connect to the database and delete the selected appointment
                    with mysql.connector.connect(
                            host="sql5.freesqldatabase.com",
                            user="sql5765928",
                            password="dhu72jpWAU",
                            database="sql5765928"
                    ) as conn:
                        cursor = conn.cursor()
                        query = "DELETE FROM appointments WHERE appointment_id = %s"
                        cursor.execute(query, (appointment_id,))
                        conn.commit()

                    messagebox.showinfo("Success", f"Appointment ID {appointment_id} deleted successfully!")
                    # Refresh the appointments table
                    tree.delete(*tree.get_children())  # Clear existing entries
                    load_appointments()  # Reload appointments

                except mysql.connector.Error as err:
                    messagebox.showerror("Database Error", f"Error: {err}")

        # Button to delete the selected appointment
        Button(self.content_frame, text="Delete Appointment", command=delete_selected_appointment, font=("Arial", 12),
               bg="red", fg="white").grid(
            row=2, column=0, columnspan=3, pady=10)
        nav_frame = tk.Frame(self.content_frame)
        nav_frame.grid(row=2, column=1, columnspan=3, pady=20)

        # Add navigation buttons
        ttk.Button(nav_frame, text="Back", command=self.manage_appointments).pack(pady=10)

    def manage_medical_records(self):
        """Manage medical records (Add, Update, Delete)."""
        for widget in self.content_frame.winfo_children():
            widget.destroy()

        manage_frame = tk.Frame(self.content_frame)
        manage_frame.pack(pady=20)
        ttk.Label = Label(manage_frame, text="Choose your choice of operation when it comes to managing Patients medical records:",
                          font=("Arial", 18))
        ttk.Label.pack()
        ttk.Button(manage_frame, text="Add Medical Record", command=self.add_medical_record, style="Main.TButton").pack(
            pady=10)
        ttk.Button(manage_frame, text="Update Medical Record", command=self.update_medical_record,
                   style="Main.TButton").pack(pady=10)
        ttk.Button(manage_frame, text="Delete Medical Record", command=self.delete_medical_record,
                   style="Main.TButton").pack(pady=10)
        self.add_navigation_buttons(back_command=lambda: self.show_admin_dashboard())

    def add_medical_record(self):
        """Function to add a new medical record with a modern and clean design."""
        # Clear previous content in the content_frame
        for widget in self.content_frame.winfo_children():
            widget.destroy()

        # Header with a modern design
        header_frame = tk.Frame(self.content_frame, bg="#4F6D7A", padx=20, pady=20)
        header_frame.pack(fill="x", pady=10)

        Label(header_frame, text="Add New Medical Record", font=("Arial", 18, "bold"), bg="#4F6D7A", fg="white").pack()

        # Create a frame for the form
        form_frame = tk.Frame(self.content_frame, bg="#f0f4fa", padx=20, pady=20)
        form_frame.pack(fill="both", expand=True)

        # Define labels and input fields for medical record details
        fields = [
            ("Patient ID:", "patient_id_entry"),
            ("Diagnosis:", "diagnosis_entry"),
            ("Treatment:", "treatment_entry"),
            ("Date of Record (YYYY-MM-DD):", "date_of_record_entry")
        ]

        input_vars = {}
        for idx, (label_text, var_name) in enumerate(fields):
            # Add labels
            label = tk.Label(form_frame, text=label_text, font=("Arial", 12), bg="#f0f4fa", fg="#34495e")
            label.grid(row=idx, column=0, padx=10, pady=5, sticky="w")

            # Add entry fields
            entry = tk.Entry(form_frame, font=("Arial", 12), width=30, bd=2, relief="groove")
            entry.grid(row=idx, column=1, padx=10, pady=5)
            input_vars[var_name] = entry

        # Function to add the medical record to the database
        def submit_medical_record():
            # Collect data from input fields
            patient_id = input_vars["patient_id_entry"].get().strip()
            diagnosis = input_vars["diagnosis_entry"].get().strip()
            treatment = input_vars["treatment_entry"].get().strip()
            date_of_record = input_vars["date_of_record_entry"].get().strip()

            # Validate inputs
            if not patient_id or not diagnosis or not treatment or not date_of_record:
                messagebox.showwarning("Input Error", "All fields are required.")
                return

            try:
                # Connect to the database to check if the patient ID exists
                with mysql.connector.connect(
                        host="sql5.freesqldatabase.com",
                        user="sql5765928",
                        password="dhu72jpWAU",
                        database="sql5765928"
                ) as conn:
                    cursor = conn.cursor()

                    # Check if the patient ID exists in the patients table
                    cursor.execute("SELECT patient_id FROM patients WHERE patient_id = %s", (patient_id,))
                    patient_exists = cursor.fetchone()

                    if not patient_exists:
                        messagebox.showerror("Invalid Patient ID",
                                             "No patient has the Patient ID you have entered. Please re-enter a valid Patient ID.")
                        return

                    # If the patient ID exists, insert the medical record
                    query = """
                        INSERT INTO medicalrecords (patient_id, diagnosis, treatment, date_of_record)
                        VALUES (%s, %s, %s, %s)
                    """
                    cursor.execute(query, (patient_id, diagnosis, treatment, date_of_record))
                    conn.commit()

                # Check if the record was successfully added
                if cursor.rowcount > 0:
                    messagebox.showinfo("Success", "Medical record added successfully!")
                    # Clear input fields after submission
                    for entry in input_vars.values():
                        entry.delete(0, "end")
                else:
                    messagebox.showwarning("Insert Failed", "Failed to add medical record.")

            except mysql.connector.Error as err:
                messagebox.showerror("Database Error", f"Error: {err}")

        # Button frame for Submit and Back buttons
        button_frame = tk.Frame(self.content_frame, bg="#f0f4fa", padx=20, pady=20)
        button_frame.pack(fill="x", pady=10)

        # Submit button with modern styling
        submit_button = tk.Button(button_frame, text="Add Record", command=submit_medical_record,
                                  font=("Arial", 12, "bold"),
                                  bg="#4CAF50", fg="white", padx=20, pady=10, relief="flat")
        submit_button.pack(side="left", padx=10)

        # Back button with modern styling
        back_button = tk.Button(button_frame, text="Back", command=self.manage_medical_records,
                                font=("Arial", 12, "bold"),
                                bg="#FF6347", fg="white", padx=20, pady=10, relief="flat")
        back_button.pack(side="right", padx=10)

        # Add hover effects for buttons
        def on_enter(event):
            event.widget.config(bg="#45a049" if event.widget == submit_button else "#FF4500")

        def on_leave(event):
            event.widget.config(bg="#4CAF50" if event.widget == submit_button else "#FF6347")

        submit_button.bind("<Enter>", on_enter)
        submit_button.bind("<Leave>", on_leave)
        back_button.bind("<Enter>", on_enter)
        back_button.bind("<Leave>", on_leave)


    def update_medical_record(self):
        """Function to update an existing medical record."""
        # Clear previous content in the content_frame
        for widget in self.content_frame.winfo_children():
            widget.destroy()

        # Instructions for the admin
        Label(self.content_frame, text="Update Medical Record", font=("Arial", 18, "bold"), fg="blue").grid(
            row=0, column=0, columnspan=3, pady=10)

        # Frame for Treeview to display existing medical records
        tree_frame = ttk.Frame(self.content_frame)
        tree_frame.grid(row=1, column=0, columnspan=3, padx=10, pady=10, sticky="nsew")

        # Define a Treeview to display medical records
        tree = ttk.Treeview(
            tree_frame,
            columns=("Record ID", "Patient ID", "Diagnosis", "Treatment", "Date of Record"),
            show="headings",
            height=8
        )
        # Define a Treeview to display medical records
        tree = ttk.Treeview(
            tree_frame,
            columns=("Record ID", "Patient ID", "Diagnosis", "Treatment", "Date of Record"),
            show="headings",
            height=15  # Increased height for more visible rows
        )
        tree.heading("Record ID", text="Record ID")
        tree.heading("Patient ID", text="Patient ID")
        tree.heading("Diagnosis", text="Diagnosis")
        tree.heading("Treatment", text="Treatment")
        tree.heading("Date of Record", text="Date of Record")

        # Adjust column widths for better visibility
        tree.column("Record ID", width=150, anchor="center", stretch=True)
        tree.column("Patient ID", width=150, anchor="center", stretch=True)
        tree.column("Diagnosis", width=200, stretch=True)
        tree.column("Treatment", width=200, stretch=True)
        tree.column("Date of Record", width=150, anchor="center", stretch=True)

        # Allow the Treeview to expand and fill the frame
        tree.pack(expand=True, fill="both")

        # Add Scrollbars
        vertical_scrollbar = ttk.Scrollbar(tree_frame, orient="vertical", command=tree.yview)
        horizontal_scrollbar = ttk.Scrollbar(tree_frame, orient="horizontal", command=tree.xview)
        tree.configure(yscrollcommand=vertical_scrollbar.set, xscrollcommand=horizontal_scrollbar.set)

        # Configure scrollbars to expand with the Treeview
        vertical_scrollbar.pack(side="right", fill="y")
        horizontal_scrollbar.pack(side="bottom", fill="x")

        # Expand the tree_frame to take available space
        tree_frame.grid(row=1, column=0, columnspan=3, padx=10, pady=10, sticky="nsew")

        # Ensure the content_frame grid expands
        self.content_frame.columnconfigure(0, weight=1)
        self.content_frame.rowconfigure(1, weight=1)

        # Function to load all medical records into the Treeview
        def load_medical_records():
            try:
                # Connect to the database to fetch all medical records
                with mysql.connector.connect(
                        host="sql5.freesqldatabase.com",
                        user="sql5765928",
                        password="dhu72jpWAU",
                        database="sql5765928"
                ) as conn:
                    cursor = conn.cursor(dictionary=True)
                    query = "SELECT * FROM medicalrecords"
                    cursor.execute(query)
                    records = cursor.fetchall()

                # Insert all records into the treeview
                for record in records:
                    tree.insert("", "end", values=(
                        record["record_id"],
                        record["patient_id"],
                        record["diagnosis"],
                        record["treatment"],
                        record["date_of_record"]
                    ))

            except mysql.connector.Error as err:
                messagebox.showerror("Database Error", f"Error: {err}")

        load_medical_records()  # Load records when the screen is shown

        # Function to update the selected medical record
        def update_selected_record():
            selected_item = tree.selection()
            if not selected_item:
                messagebox.showwarning("Selection Error", "Please select a medical record to update.")
                return

            record_id = tree.item(selected_item[0], "values")[0]

            # Fetch the details of the selected record
            try:
                with mysql.connector.connect(
                        host="sql5.freesqldatabase.com",
                        user="sql5765928",
                        password="dhu72jpWAU",
                        database="sql5765928"
                ) as conn:
                    cursor = conn.cursor(dictionary=True)
                    query = "SELECT * FROM medicalrecords WHERE record_id = %s"
                    cursor.execute(query, (record_id,))
                    record = cursor.fetchone()

                if not record:
                    messagebox.showwarning("Record Not Found", "No record found with the selected ID.")
                    return

                # Create a form to update the selected record
                record_form_frame = ttk.Frame(self.content_frame)
                record_form_frame.grid(row=2, column=0, columnspan=3, pady=10)

                Label(record_form_frame, text="Diagnosis:", font=("Arial", 12)).grid(row=0, column=0, padx=5, pady=5)
                diagnosis_entry = Entry(record_form_frame, font=("Arial", 12))
                diagnosis_entry.insert(0, record["diagnosis"])
                diagnosis_entry.grid(row=0, column=1, padx=5, pady=5)

                Label(record_form_frame, text="Treatment:", font=("Arial", 12)).grid(row=1, column=0, padx=5, pady=5)
                treatment_entry = Entry(record_form_frame, font=("Arial", 12))
                treatment_entry.insert(0, record["treatment"])
                treatment_entry.grid(row=1, column=1, padx=5, pady=5)

                Label(record_form_frame, text="Date of Record:", font=("Arial", 12)).grid(row=2, column=0, padx=5,
                                                                                          pady=5)
                date_of_record_entry = Entry(record_form_frame, font=("Arial", 12))
                date_of_record_entry.insert(0, record["date_of_record"])
                date_of_record_entry.grid(row=2, column=1, padx=5, pady=5)

                # Function to save the updated record
                def save_updated_record():
                    updated_diagnosis = diagnosis_entry.get()
                    updated_treatment = treatment_entry.get()
                    updated_date_of_record = date_of_record_entry.get()

                    if not updated_diagnosis or not updated_treatment or not updated_date_of_record:
                        messagebox.showwarning("Input Error", "Please fill in all fields.")
                        return

                    try:
                        with mysql.connector.connect(
                                host="sql5.freesqldatabase.com",
                                user="sql5765928",
                                password="dhu72jpWAU",
                                database="sql5765928"
                        ) as conn:
                            cursor = conn.cursor()
                            query = """
                                UPDATE medicalrecords
                                SET diagnosis = %s, treatment = %s, date_of_record = %s
                                WHERE record_id = %s
                            """
                            cursor.execute(query,
                                           (updated_diagnosis, updated_treatment, updated_date_of_record, record_id))
                            conn.commit()

                        messagebox.showinfo("Success", "Medical record updated successfully!")
                        # Refresh the treeview
                        tree.delete(*tree.get_children())  # Clear the current records
                        load_medical_records()  # Reload the updated list

                    except mysql.connector.Error as err:
                        messagebox.showerror("Database Error", f"Error: {err}")

                # Button to save the updated record
                Button(record_form_frame, text="Save Changes", command=save_updated_record, font=("Arial", 12),
                       bg="green", fg="white").grid(
                    row=3, column=0, columnspan=2, pady=10)

            except mysql.connector.Error as err:
                messagebox.showerror("Database Error", f"Error: {err}")

        # Button to update the selected medical record
        Button(self.content_frame, text="Update Medical Record", command=update_selected_record, font=("Arial", 12),
               bg="blue", fg="white").grid(
            row=3, column=0, columnspan=3, pady=10)
        nav_frame = tk.Frame(self.content_frame)
        nav_frame.grid(row=3, column=2, columnspan=3, pady=20)

        # Add navigation buttons
        ttk.Button(nav_frame, text="Back", command=self.manage_medical_records).pack(pady=10)

    def delete_medical_record(self):
        """Function to delete a medical record."""
        # Clear previous content in the content_frame
        for widget in self.content_frame.winfo_children():
            widget.destroy()

        # Instructions for the admin
        Label(self.content_frame, text="Delete Medical Record", font=("Arial", 18, "bold"), fg="blue").grid(
            row=0, column=0, columnspan=3, pady=10)

        # Frame for the Treeview to display all medical records
        tree_frame = ttk.Frame(self.content_frame)
        tree_frame.grid(row=1, column=0, columnspan=3, padx=10, pady=10, sticky="nsew")

        # Define a Treeview to display all medical records
        tree = ttk.Treeview(
            tree_frame,
            columns=("Record ID", "Patient ID", "Diagnosis", "Treatment", "Date of Record"),
            show="headings",
            height=8
        )
        tree.heading("Record ID", text="Record ID")
        tree.heading("Patient ID", text="Patient ID")
        tree.heading("Diagnosis", text="Diagnosis")
        tree.heading("Treatment", text="Treatment")
        tree.heading("Date of Record", text="Date of Record")

        tree.column("Record ID", width=100, anchor="center", stretch=True)
        tree.column("Patient ID", width=100, anchor="center", stretch=True)
        tree.column("Diagnosis", width=150, anchor="center", stretch=True)
        tree.column("Treatment", width=150, anchor="center", stretch=True)
        tree.column("Date of Record", width=150, anchor="center", stretch=True)

        tree.pack(expand=True, fill="both")

        # Add Scrollbars
        vertical_scrollbar = ttk.Scrollbar(tree_frame, orient="vertical", command=tree.yview)
        horizontal_scrollbar = ttk.Scrollbar(tree_frame, orient="horizontal", command=tree.xview)
        tree.configure(yscrollcommand=vertical_scrollbar.set, xscrollcommand=horizontal_scrollbar.set)

        vertical_scrollbar.pack(side="right", fill="y")
        horizontal_scrollbar.pack(side="bottom", fill="x")

        # Function to load all medical records into the Treeview
        def load_medical_records():
            try:
                # Connect to the database to fetch all medical records
                with mysql.connector.connect(
                        host="sql5.freesqldatabase.com",
                        user="sql5765928",
                        password="dhu72jpWAU",
                        database="sql5765928"
                ) as conn:
                    cursor = conn.cursor(dictionary=True)
                    query = "SELECT * FROM medicalrecords"
                    cursor.execute(query)
                    records = cursor.fetchall()

                # Insert all medical records into the treeview
                for record in records:
                    tree.insert("", "end", values=(
                        record["record_id"],
                        record["patient_id"],
                        record["diagnosis"],
                        record["treatment"],
                        record["date_of_record"]
                    ))

            except mysql.connector.Error as err:
                messagebox.showerror("Database Error", f"Error: {err}")

        load_medical_records()  # Load records when the screen is shown

        # Function to delete the selected medical record
        def delete_selected_record():
            selected_item = tree.selection()
            if not selected_item:
                messagebox.showwarning("Selection Error", "Please select a medical record to delete.")
                return

            record_id = tree.item(selected_item[0], "values")[0]

            # Ask for confirmation before deleting
            confirm = messagebox.askyesno("Confirm Deletion", f"Are you sure you want to delete record ID {record_id}?")
            if confirm:
                try:
                    # Connect to the database and delete the selected record
                    with mysql.connector.connect(
                            host="sql5.freesqldatabase.com",
                            user="sql5765928",
                            password="dhu72jpWAU",
                            database="sql5765928"
                    ) as conn:
                        cursor = conn.cursor()
                        query = "DELETE FROM medicalrecords WHERE record_id = %s"
                        cursor.execute(query, (record_id,))
                        conn.commit()

                    messagebox.showinfo("Success", f"Record ID {record_id} deleted successfully!")
                    # Refresh the medical records table
                    tree.delete(*tree.get_children())  # Clear existing entries
                    load_medical_records()  # Reload medical records

                except mysql.connector.Error as err:
                    messagebox.showerror("Database Error", f"Error: {err}")

        # Button to delete the selected medical record
        Button(self.content_frame, text="Delete Record", command=delete_selected_record, font=("Arial", 12), bg="red",
               fg="white").grid(
            row=2, column=0, columnspan=3, pady=10)
        nav_frame = tk.Frame(self.content_frame)
        nav_frame.grid(row=2, column=1, columnspan=3, pady=20)

        # Add navigation buttons
        ttk.Button(nav_frame, text="Back", command=self.manage_medical_records).pack(pady=10)

    def show_patient_dashboard(self, patient_data):
        """Display the patient's dashboard."""
        for widget in self.content_frame.winfo_children():
            widget.destroy()
        if isinstance(patient_data, int):  # If patient_data is patient_id, fetch full data
            try:
                with mysql.connector.connect(
                        host="sql5.freesqldatabase.com",
                        user="sql5765928",
                        password="dhu72jpWAU",
                        database="sql5765928"
                ) as conn:
                    cursor = conn.cursor(dictionary=True)
                    query = "SELECT * FROM patients WHERE patient_id = %s"
                    cursor.execute(query, (patient_data,))
                    patient_data = cursor.fetchone()

                if not patient_data:
                    messagebox.showerror("Error", "Patient data not found.")
                    return

            except mysql.connector.Error as err:
                messagebox.showerror("Database Error", f"Error: {err}")
                return

        tk.Label(
            self.content_frame,
            text=f"Welcome, {patient_data['first_name']} {patient_data['last_name']}!",
            font=("Helvetica", 16, "bold"),
            bg="#f0f4fa",
        ).pack(pady=10)

        options_frame = tk.Frame(self.content_frame, bg="#f0f4fa")
        options_frame.pack(pady=20)

        ttk.Button(
            options_frame,
            text="View Medical Records",
            command=lambda: self.view_medical_records(patient_data['patient_id']),
            style="Main.TButton",
        ).pack(pady=10)

        ttk.Button(
            options_frame,
            text="View Appointments",
            command=lambda: self.view_appointments(patient_data['patient_id']),
            style="Main.TButton",
        ).pack(pady=10)

        ttk.Button(
            options_frame,
            text="Update Profile",
            command=lambda: self.update_profile(patient_data),
            style="Main.TButton",
        ).pack(pady=10)

        ttk.Button(
            options_frame,
            text="Logout",
            command=self.show_home,
            style="Main.TButton",
        ).pack(pady=10)

    def show_about(self):
        # Clear existing content
        for widget in self.content_frame.winfo_children():
            widget.destroy()

        # Set custom background color
        self.content_frame.configure(bg="#e8f1f8")  # Light blue-gray

        # Add a title with improved styling
        tk.Label(
            self.content_frame,
            text="About Us",
            font=("Helvetica", 20, "bold"),
            fg="#2c3e50",  # Dark Blue
            bg="#e8f1f8",
            pady=10,
        ).pack(pady=10)

        # Create a frame for the content box
        content_box = tk.Frame(self.content_frame, bg="#ffffff", bd=2, relief="groove")
        content_box.pack(pady=20, padx=20, fill="both", expand=True)

        # Add an icon/image (using Pillow to handle .jpg file)
        try:
            original_image = Image.open("background.jpg")  # Open the image file
            resized_image = original_image.resize((200, 200))  # Resize the image as needed
            icon = ImageTk.PhotoImage(resized_image)  # Convert to PhotoImage for Tkinter
            tk.Label(content_box, image=icon, bg="#ffffff").pack(pady=10)
            self.content_frame.icon = icon  # Prevent garbage collection
        except Exception as e:
            print(f"Error loading image: {e}")

        # Add hospital information in a nicely formatted way
        info_text = (
            "At J.F.K Hospital, your health is our priority.\n\n"
            "Feel free to reach out to us anytime using the information below:\n\n"
            "📞 WhatsApp: +250790878665\n"
            "📧 Email: hyallison5050@gmail.com\n"
            "📍 Location: Kimonyi, Musanze District, Rwanda"
        )
        tk.Label(
            content_box,
            text=info_text,
            font=("Helvetica", 14),
            fg="#34495e",  # Dark Gray
            bg="#ffffff",
            justify="left",
            wraplength=500,
        ).pack(pady=20)

        # Add navigation buttons
        self.add_navigation_buttons(back_command=self.show_home, show_logout=True)




    import tkinter as tk
    from tkinter import ttk, messagebox
    import mysql.connector

    class PatientApp:
        def __init__(self, root, content_frame):
            self.root = root
            self.content_frame = content_frame

    def view_medical_records(self, patient_id):
        """Display the patient's medical records in a tabular format."""
        # Clear the previous widgets in content_frame
        for widget in self.content_frame.winfo_children():
            widget.destroy()

        # Title Label
        tk.Label(
            self.content_frame,
            text="Your Medical Records",
            font=("Helvetica", 16, "bold"),
            bg="#f0f4fa",
        ).pack(pady=10)

        # Frame to hold the table and scrollbars
        records_frame = tk.Frame(self.content_frame, bg="#f0f4fa")
        records_frame.pack(pady=10, fill="both", expand=True)

        # Create a frame to hold the Treeview and scrollbars
        table_frame = tk.Frame(records_frame)
        table_frame.pack(fill="both", expand=True)

        # Add vertical and horizontal scrollbars
        v_scrollbar = ttk.Scrollbar(table_frame, orient="vertical")
        v_scrollbar.pack(side="right", fill="y")
        h_scrollbar = ttk.Scrollbar(table_frame, orient="horizontal")
        h_scrollbar.pack(side="bottom", fill="x")

        # Create the Treeview for displaying the table
        tree = ttk.Treeview(
            table_frame,
            columns=("Diagnosis", "Treatment", "Date of Record"),
            show="headings",
            yscrollcommand=v_scrollbar.set,
            xscrollcommand=h_scrollbar.set
        )

        # Configure the scrollbars to work with the Treeview
        v_scrollbar.config(command=tree.yview)
        h_scrollbar.config(command=tree.xview)

        # Define the column headings
        tree.heading("Diagnosis", text="Diagnosis")
        tree.heading("Treatment", text="Treatment")
        tree.heading("Date of Record", text="Date of Record")

        # Define the column width and anchor
        tree.column("Diagnosis", width=300, anchor="center")  # Increased width
        tree.column("Treatment", width=300, anchor="center")  # Increased width
        tree.column("Date of Record", width=200, anchor="center")  # Increased width

        # Styling the Treeview
        tree.tag_configure("oddrow", background="#f9f9f9")
        tree.tag_configure("evenrow", background="#ffffff")

        # Place the Treeview widget in the frame
        tree.pack(fill="both", expand=True)

        try:
            # Using a context manager to automatically close the connection
            with mysql.connector.connect(
                    host="sql5.freesqldatabase.com",
                    user="sql5765928",
                    password="dhu72jpWAU",
                    database="sql5765928"
            ) as conn:
                cursor = conn.cursor()
                query = "SELECT patient_id, diagnosis, treatment, date_of_record FROM medicalrecords WHERE patient_id = %s"
                cursor.execute(query, (patient_id,))
                records = cursor.fetchall()

            # Check if there are medical records and add them to the table
            if records:
                for index, record in enumerate(records):
                    row_tag = "oddrow" if index % 2 == 0 else "evenrow"
                    tree.insert("", "end", values=(record[1], record[2], record[3]), tags=(row_tag,))
            else:
                # No records found
                tk.Label(records_frame, text="No records found.", bg="#f0f4fa").pack(pady=5)

        except mysql.connector.Error as err:
            messagebox.showerror("Database Error", f"Error: {err}")

        self.add_navigation_buttons(back_command=lambda: self.show_patient_dashboard(patient_id))

    def view_appointments(self, patient_id):
        """Display the patient's appointments in a tabular format."""
        # Clear the previous widgets in content_frame
        for widget in self.content_frame.winfo_children():
            widget.destroy()

        # Title Label
        tk.Label(
            self.content_frame,
            text="Your Appointments",
            font=("Helvetica", 16, "bold"),
            bg="#f0f4fa",
        ).pack(pady=10)

        # Frame to hold the table with scrollbars
        appointments_frame = tk.Frame(self.content_frame, bg="#f0f4fa")
        appointments_frame.pack(pady=10, fill="both", expand=True)

        # Create a Scrollbar for vertical scrolling
        vertical_scrollbar = tk.Scrollbar(appointments_frame, orient="vertical")
        vertical_scrollbar.pack(side="right", fill="y")

        # Create a Scrollbar for horizontal scrolling
        horizontal_scrollbar = tk.Scrollbar(appointments_frame, orient="horizontal")
        horizontal_scrollbar.pack(side="bottom", fill="x")

        # Create the Treeview for displaying the table
        tree = ttk.Treeview(
            appointments_frame,
            columns=("Doctor ID", "Appointment Date", "Status", "Purpose"),
            show="headings",
            yscrollcommand=vertical_scrollbar.set,
            xscrollcommand=horizontal_scrollbar.set,
        )

        # Attach scrollbars to Treeview
        vertical_scrollbar.config(command=tree.yview)
        horizontal_scrollbar.config(command=tree.xview)

        # Define the column headings
        tree.heading("Doctor ID", text="Doctor ID")
        tree.heading("Appointment Date", text="Appointment Date")
        tree.heading("Status", text="Status")
        tree.heading("Purpose", text="Purpose")

        # Define the column width and anchor
        tree.column("Doctor ID", width=150, anchor="center")
        tree.column("Appointment Date", width=200, anchor="center")
        tree.column("Status", width=150, anchor="center")
        tree.column("Purpose", width=250, anchor="center")

        # Styling the Treeview
        tree.tag_configure("oddrow", background="#f9f9f9")
        tree.tag_configure("evenrow", background="#ffffff")

        # Place the Treeview widget in the frame
        tree.pack(fill="both", expand=True)

        try:
            # Using a context manager to automatically close the connection
            with mysql.connector.connect(
                    host="sql5.freesqldatabase.com",
                    user="sql5765928",
                    password="dhu72jpWAU",
                    database="sql5765928"
            ) as conn:
                cursor = conn.cursor()
                query = "SELECT patient_id, doctor_id, appointment_date, status, purpose FROM appointments WHERE patient_id = %s"
                cursor.execute(query, (patient_id,))
                appointments = cursor.fetchall()

            # Check if there are appointments and add them to the table
            if appointments:
                for index, appointment in enumerate(appointments):
                    row_tag = "oddrow" if index % 2 == 0 else "evenrow"
                    tree.insert("", "end", values=(appointment[1], appointment[2], appointment[3], appointment[4]),
                                tags=(row_tag,))
            else:
                # No appointments found
                tk.Label(appointments_frame, text="No appointments found.", bg="#f0f4fa").pack(pady=5)

        except mysql.connector.Error as err:
            messagebox.showerror("Database Error", f"Error: {err}")
        self.add_navigation_buttons(lambda: self.show_patient_dashboard(patient_id))

    def update_profile(self, patient_data):
        """Allow the patient to update their profile and view their current information."""
        # Clear the previous widgets in content_frame
        for widget in self.content_frame.winfo_children():
            widget.destroy()

        # Title Label
        tk.Label(
            self.content_frame,
            text="Update Profile Information",
            font=("Helvetica", 16, "bold"),
            bg="#f0f4fa",
        ).pack(pady=10)

        # Frame to display current patient information
        info_frame = tk.Frame(self.content_frame, bg="#f0f4fa")
        info_frame.pack(pady=10)

        # Create the Treeview for displaying the current profile information
        tree = ttk.Treeview(info_frame, columns=("Field", "Value"), show="headings", height=8)

        # Define the column headings
        tree.heading("Field", text="Field")
        tree.heading("Value", text="Value")

        # Define the column width and anchor
        tree.column("Field", width=150, anchor="center")
        tree.column("Value", width=250, anchor="w")

        # Styling the Treeview
        tree.tag_configure("oddrow", background="#f9f9f9")
        tree.tag_configure("evenrow", background="#ffffff")

        # Place the Treeview widget in the frame
        tree.pack(fill="both", expand=True)

        # Insert the current profile data into the table
        current_info = [
            ("First Name", patient_data.get("first_name", "N/A")),
            ("Last Name", patient_data.get("last_name", "N/A")),
            ("Gender", patient_data.get("gender", "N/A")),
            ("Email", patient_data.get("email", "N/A")),
            ("Contact", patient_data.get("contact", "N/A")),
            ("Address", patient_data.get("address", "N/A")),
        ]

        for index, info in enumerate(current_info):
            row_tag = "oddrow" if index % 2 == 0 else "evenrow"
            tree.insert("", "end", values=info, tags=(row_tag,))

        # Frame to hold the form for updating profile
        form_frame = tk.Frame(self.content_frame, bg="#f0f4fa")
        form_frame.pack(pady=20)

        # First Name
        tk.Label(form_frame, text="First Name:", bg="#f0f4fa").grid(row=0, column=0, pady=5, sticky="w")
        first_name_entry = ttk.Entry(form_frame, width=30)
        first_name_entry.insert(0, patient_data["first_name"])
        first_name_entry.grid(row=0, column=1, pady=5)

        # Last Name
        tk.Label(form_frame, text="Last Name:", bg="#f0f4fa").grid(row=1, column=0, pady=5, sticky="w")
        last_name_entry = ttk.Entry(form_frame, width=30)
        last_name_entry.insert(0, patient_data["last_name"])
        last_name_entry.grid(row=1, column=1, pady=5)

        # Gender
        tk.Label(form_frame, text="Gender:", bg="#f0f4fa").grid(row=2, column=0, pady=5, sticky="w")
        gender_entry = ttk.Entry(form_frame, width=30)
        gender_entry.insert(0, patient_data["gender"])
        gender_entry.grid(row=2, column=1, pady=5)

        # Email
        tk.Label(form_frame, text="Email:", bg="#f0f4fa").grid(row=3, column=0, pady=5, sticky="w")
        email_entry = ttk.Entry(form_frame, width=30)
        email_entry.insert(0, patient_data["email"])
        email_entry.grid(row=3, column=1, pady=5)

        # Contact
        tk.Label(form_frame, text="Contact:", bg="#f0f4fa").grid(row=4, column=0, pady=5, sticky="w")
        contact_entry = ttk.Entry(form_frame, width=30)
        contact_entry.insert(0, patient_data["contact"])
        contact_entry.grid(row=4, column=1, pady=5)

        # Address
        tk.Label(form_frame, text="Address:", bg="#f0f4fa").grid(row=5, column=0, pady=5, sticky="w")
        address_entry = ttk.Entry(form_frame, width=30)
        address_entry.insert(0, patient_data["address"])
        address_entry.grid(row=5, column=1, pady=5)

        # Password
        tk.Label(form_frame, text="Password:", bg="#f0f4fa").grid(row=6, column=0, pady=5, sticky="w")
        password_entry = ttk.Entry(form_frame, show="*", width=30)
        password_entry.grid(row=6, column=1, pady=5)

        # Save Profile Logic
        def save_profile():
            first_name = first_name_entry.get().strip()
            last_name = last_name_entry.get().strip()
            gender = gender_entry.get().strip()
            email = email_entry.get().strip()
            contact = contact_entry.get().strip()
            address = address_entry.get().strip()
            password = password_entry.get().strip()

            # Validate fields are not empty
            if not (first_name and last_name and gender and email and contact and address and password):
                messagebox.showerror("Error", "All fields are required!")
                return

            try:
                conn = mysql.connector.connect(
                    host="sql5.freesqldatabase.com",
                    user="sql5765928",
                    password="dhu72jpWAU",
                    database="sql5765928"
                )
                cursor = conn.cursor()
                query = """UPDATE patients 
                           SET first_name = %s, last_name = %s, gender = %s, email = %s, 
                               contact = %s, address = %s, password = %s 
                           WHERE patient_id = %s"""
                cursor.execute(query, (
                first_name, last_name, gender, email, contact, address, password, patient_data["patient_id"]))
                conn.commit()
                cursor.close()
                conn.close()

                # Update the patient_data dictionary with new info
                patient_data.update({
                    "first_name": first_name,
                    "last_name": last_name,
                    "gender": gender,
                    "email": email,
                    "contact": contact,
                    "address": address,
                    "password": password,
                })

                messagebox.showinfo("Success", "Profile updated successfully!")
                self.show_patient_dashboard(patient_data)

            except mysql.connector.Error as err:
                messagebox.showerror("Database Error", f"Error: {err}")

        ttk.Button(
            form_frame,
            text="Save Changes",
            command=save_profile,
            style="Main.TButton",
        ).grid(row=7, column=0, columnspan=2, pady=20)

        # Navigation Buttons
        self.add_navigation_buttons(back_command=lambda: self.show_patient_dashboard(patient_data))


if __name__ == "__main__":
    root = ThemedTk(theme="adapta")
    app = PatientManagementApp(root)
    root.mainloop()